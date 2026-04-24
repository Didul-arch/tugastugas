"""
Generate LKP 6 DOCX report — format beralur (kode → penjelasan → screenshot).
Run from pertemuan6/ folder.
"""

import cv2
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
os.chdir(SCRIPT_DIR)

OUT_DIR = 'report_imgs'
os.makedirs(OUT_DIR, exist_ok=True)

# ── Helper functions ──
def add_code_block(doc, code_text):
    """Add a formatted code block (gray background via font styling)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)

def add_explanation(doc, text):
    """Add explanation paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)

def add_img_centered(doc, path, width=Inches(5)):
    """Add centered image."""
    doc.add_picture(path, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════
#  GENERATE FIGURES
# ═══════════════════════════════════════════

img = cv2.imread('tomato.jpeg')
rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
h, s, v = cv2.split(hsv)

# Fig 1: Original
plt.figure(figsize=(4, 4))
plt.imshow(rgb); plt.title('Original'); plt.axis('off')
plt.savefig(f'{OUT_DIR}/01_original.png', dpi=150, bbox_inches='tight'); plt.close()

# Fig 2: HSV channels + histogram
fig, axes = plt.subplots(2, 3, figsize=(14, 6))
for i, (ch, name) in enumerate(zip([h, s, v], ['H', 'S', 'V'])):
    axes[0, i].imshow(ch, cmap='gray'); axes[0, i].set_title(f'Channel {name}'); axes[0, i].axis('off')
    axes[1, i].hist(ch.ravel(), bins=256, range=[0, 256], color='black'); axes[1, i].set_title(f'Histogram {name}')
plt.tight_layout(); plt.savefig(f'{OUT_DIR}/02_hsv_histogram.png', dpi=150, bbox_inches='tight'); plt.close()

# Fig 3: Pipeline
mask1 = cv2.inRange(hsv, np.array([0, 27, 40]), np.array([15, 255, 255]))
mask2 = cv2.inRange(hsv, np.array([165, 27, 40]), np.array([179, 255, 255]))
binary = cv2.bitwise_or(mask1, mask2)
k_erosi = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
k_dilasi = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
erosi = cv2.erode(binary, k_erosi)
dilasi = cv2.dilate(erosi, k_dilasi)
hasil = cv2.bitwise_and(rgb, rgb, mask=dilasi)

fig, axes = plt.subplots(1, 4, figsize=(16, 4))
for i, (im, t) in enumerate(zip([binary, erosi, dilasi, hasil], ['Binary', 'Erosi 3x3', 'Dilasi 5x5', 'Hasil Segmentasi'])):
    axes[i].imshow(im, cmap='gray' if i < 3 else None); axes[i].set_title(t); axes[i].axis('off')
plt.tight_layout(); plt.savefig(f'{OUT_DIR}/03_pipeline.png', dpi=150, bbox_inches='tight'); plt.close()

# Fig 4: Variation grid
configs = [(3,3,1,1),(3,5,1,1),(5,5,1,1),(5,7,1,1),(3,5,2,1),(3,5,1,2)]
fig, axes = plt.subplots(len(configs), 4, figsize=(16, 3.5 * len(configs)))
for row, (ke_s, kd_s, ie, id_) in enumerate(configs):
    ke = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (ke_s, ke_s))
    kd = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (kd_s, kd_s))
    er = cv2.erode(binary, ke, iterations=ie)
    dl = cv2.dilate(er, kd, iterations=id_)
    seg = cv2.bitwise_and(rgb, rgb, mask=dl)
    for col, (im, ttl) in enumerate(zip(
        [binary, er, dl, seg],
        ['Binary', f'Erosi {ke_s}x{ke_s} i={ie}', f'Dilasi {kd_s}x{kd_s} i={id_}', 'Hasil']
    )):
        axes[row, col].imshow(im, cmap='gray' if col < 3 else None)
        axes[row, col].set_title(ttl, fontsize=10); axes[row, col].axis('off')
plt.tight_layout(); plt.savefig(f'{OUT_DIR}/04_variasi.png', dpi=150, bbox_inches='tight'); plt.close()

print("Figures saved.")

# ═══════════════════════════════════════════
#  BUILD DOCX — FORMAT BERALUR
# ═══════════════════════════════════════════

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ── Title ──
title = doc.add_heading('LKP 6 — Morfologi', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Syafiq Syadidul Azmi — G6401231075')
run.font.size = Pt(12); run.font.name = 'Times New Roman'

# ═══════════════════════════════════════════
#  LANGKAH 1: LOAD CITRA
# ═══════════════════════════════════════════
doc.add_heading('Langkah 1 — Load Citra', level=2)

add_code_block(doc, """\
import cv2
import matplotlib.pyplot as plt
import numpy as np

img = cv2.imread('tomato.jpeg')
rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
h, s, v = cv2.split(hsv)""")

add_explanation(doc,
    'Import library yang dibutuhkan: cv2 (OpenCV) untuk pengolahan citra, '
    'matplotlib untuk visualisasi, dan numpy untuk operasi array.')
add_explanation(doc,
    'cv2.imread() membaca file gambar tomato.jpeg dalam format BGR. '
    'Kemudian dikonversi ke RGB (untuk ditampilkan dengan matplotlib) dan ke HSV (untuk segmentasi). '
    'cv2.split() memisahkan citra HSV menjadi 3 channel: H (Hue/warna), S (Saturation/kejenuhan), V (Value/kecerahan).')

add_img_centered(doc, f'{OUT_DIR}/01_original.png', Inches(3))

# ═══════════════════════════════════════════
#  LANGKAH 2: ANALISIS CHANNEL HSV
# ═══════════════════════════════════════════
doc.add_heading('Langkah 2 — Analisis Channel HSV & Histogram', level=2)

add_code_block(doc, """\
fig, axes = plt.subplots(2, 3, figsize=(15, 7))

for i, (ch, name) in enumerate(zip([h, s, v], ['H', 'S', 'V'])):
    axes[0, i].imshow(ch, cmap='gray')
    axes[0, i].set_title(f'Channel {name}')
    axes[1, i].hist(ch.ravel(), bins=256, range=[0, 256], color='black')
    axes[1, i].set_title(f'Histogram {name}')""")

add_explanation(doc,
    'Menampilkan visualisasi tiap channel (H, S, V) beserta histogramnya dalam grid 2x3. '
    'Baris atas menampilkan citra grayscale tiap channel, baris bawah menampilkan distribusi intensitas pikselnya.')
add_explanation(doc,
    'Pemilihan color space HSV karena channel H langsung merepresentasikan jenis warna — '
    'cocok untuk segmentasi objek berwarna spesifik seperti tomat merah. '
    'Dari histogram H, terlihat spike besar di sekitar 0-15 yang merupakan piksel merah tomat. '
    'Histogram S menunjukkan tomat memiliki saturasi tinggi (jenuh), background putih saturasi rendah. '
    'Histogram V menunjukkan keduanya sama-sama terang.')

add_img_centered(doc, f'{OUT_DIR}/02_hsv_histogram.png', Inches(6))

# ═══════════════════════════════════════════
#  LANGKAH 3: THRESHOLDING
# ═══════════════════════════════════════════
doc.add_heading('Langkah 3 — Thresholding HSV', level=2)

add_code_block(doc, """\
mask1 = cv2.inRange(hsv, np.array([0, 27, 40]), np.array([15, 255, 255]))
mask2 = cv2.inRange(hsv, np.array([165, 27, 40]), np.array([179, 255, 255]))
binary = cv2.bitwise_or(mask1, mask2)""")

add_explanation(doc,
    'cv2.inRange() membuat mask biner: piksel yang berada dalam rentang [lower, upper] bernilai 255 (putih), sisanya 0 (hitam).')
add_explanation(doc,
    'Digunakan 2 range karena channel H bersifat melingkar — warna merah ada di dua ujung: '
    '0-15 (merah-oranye) dan 165-179 (merah-magenta). '
    'mask1 menangkap merah di ujung kiri, mask2 di ujung kanan. '
    'Keduanya digabung dengan bitwise_or sehingga semua piksel merah tercakup.')
add_explanation(doc,
    'S minimum = 27 untuk membuang piksel yang terlalu pudar (highlight/background putih). '
    'V minimum = 40 untuk membuang piksel terlalu gelap (bayangan).')

# ═══════════════════════════════════════════
#  LANGKAH 4: MORFOLOGI (EROSI + DILASI)
# ═══════════════════════════════════════════
doc.add_heading('Langkah 4 — Morfologi: Erosi dan Dilasi', level=2)

add_code_block(doc, """\
k_erosi = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
k_dilasi = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
erosi = cv2.erode(binary, k_erosi)
dilasi = cv2.dilate(erosi, k_dilasi)
hasil = cv2.bitwise_and(rgb, rgb, mask=dilasi)""")

add_explanation(doc,
    'cv2.getStructuringElement() membuat kernel (structuring element) berbentuk ellipse. '
    'Kernel erosi berukuran 3x3 — cukup kecil untuk menghilangkan noise tanpa mengikis objek terlalu banyak.')
add_explanation(doc,
    'cv2.erode() mengecilkan area putih pada citra biner — noise kecil di tepi objek hilang. '
    'cv2.dilate() melebarkan kembali area putih — mengembalikan bentuk yang terkikis erosi '
    'sekaligus menutup celah kecil di dalam objek. Kernel dilasi 5x5 sengaja lebih besar dari erosi '
    'agar bentuk objek pulih sepenuhnya.')
add_explanation(doc,
    'cv2.bitwise_and() menerapkan mask hasil dilasi ke citra RGB original — '
    'piksel yang sesuai mask ditampilkan, sisanya hitam. Ini adalah hasil segmentasi akhir.')

add_img_centered(doc, f'{OUT_DIR}/03_pipeline.png', Inches(6.5))

# ═══════════════════════════════════════════
#  LANGKAH 5: VARIASI KERNEL
# ═══════════════════════════════════════════
doc.add_heading('Langkah 5 — Variasi Ukuran Kernel dan Iterasi', level=2)

add_code_block(doc, """\
configs = [
    (3, 3, 1, 1),   # erosi 3x3, dilasi 3x3
    (3, 5, 1, 1),   # erosi 3x3, dilasi 5x5  ← pilihan akhir
    (5, 5, 1, 1),   # erosi 5x5, dilasi 5x5
    (5, 7, 1, 1),   # erosi 5x5, dilasi 7x7
    (3, 5, 2, 1),   # erosi 3x3 2 iter, dilasi 5x5
    (3, 5, 1, 2),   # erosi 3x3, dilasi 5x5 2 iter
]

for row, (ke_s, kd_s, ie, id_) in enumerate(configs):
    ke = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (ke_s, ke_s))
    kd = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (kd_s, kd_s))
    er = cv2.erode(binary, ke, iterations=ie)
    dl = cv2.dilate(er, kd, iterations=id_)
    seg = cv2.bitwise_and(rgb, rgb, mask=dl)""")

add_explanation(doc,
    'configs berisi 6 kombinasi parameter: (ukuran_kernel_erosi, ukuran_kernel_dilasi, iterasi_erosi, iterasi_dilasi). '
    'Loop menguji tiap kombinasi secara berurutan untuk membandingkan hasilnya.')
add_explanation(doc,
    'getStructuringElement() membuat kernel ellipse sesuai ukuran yang ditentukan. '
    'erode() dan dilate() menerima parameter iterations untuk mengulang operasi beberapa kali.')

add_img_centered(doc, f'{OUT_DIR}/04_variasi.png', Inches(6.5))

doc.add_heading('Pengaruh Perubahan Kernel', level=3)

doc.add_paragraph(
    'Kernel kecil (3x3): detail tepi terjaga, tapi noise kecil masih bisa tersisa.',
    style='List Bullet')
doc.add_paragraph(
    'Kernel besar (5x5, 7x7): noise lebih bersih, tapi tepi objek terkikis/melebar berlebihan.',
    style='List Bullet')
doc.add_paragraph(
    'Iterasi erosi tinggi: objek makin menyusut, bagian tipis bisa hilang.',
    style='List Bullet')
doc.add_paragraph(
    'Iterasi dilasi tinggi: objek membesar, bisa melebar melewati batas asli.',
    style='List Bullet')

doc.add_heading('Kesimpulan — Alasan Pemilihan Kernel Akhir', level=3)
add_explanation(doc,
    'Dipilih erosi 3x3 (1 iterasi) + dilasi 5x5 (1 iterasi) karena: '
    'erosi 3x3 cukup untuk menghilangkan noise kecil tanpa mengikis objek; '
    'dilasi 5x5 mengembalikan bentuk yang terkikis sekaligus menutup celah; '
    'masing-masing 1 iterasi sudah memberikan hasil bersih dan natural.')

# ── Save ──
output_path = 'LKP6_G6401231075_Syafiq_Syadidul_Azmi.docx'
doc.save(output_path)
print(f"\nDOCX saved: {output_path}")
